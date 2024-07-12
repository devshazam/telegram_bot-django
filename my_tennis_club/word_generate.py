from docx import Document
from docx.shared import Inches, Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import io
from members.models import Clients
from num2words import num2words
import datetime
import locale
locale.setlocale(locale.LC_ALL, "ru") 


def word_generate(clientArray, dayNumber):

    listOfClients = Clients.objects.get(id=clientArray[0])
    print(clientArray)
    print(listOfClients.name)
    document = Document()
    # задаем стиль текста по умолчанию
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)


    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('Продавец: Мовсисян Э.Н.')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('Адрес: 400096. ул. Удмуртская д.99 кв. 169')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('ИНН: 344790398647')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('Расчетный счет: 40802810514100011946')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('Кор. счет: 30101810000000000201')
    sentence.font.size = Pt(8)
    
    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('Банк:  ПАО АКБ "АВАНГАРД" г. МОСКВА')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('ИНН: 7702021163')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('КПП: 0')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('БИК: 044525201')
    sentence.font.size = Pt(8)

    p = document.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = p.add_run('БИК: 044525201')
    sentence.font.size = Pt(8)

    p = document.add_paragraph('')

 
    todayDae = datetime.datetime.today().strftime("%d.%b.%y")
    
    paragraph = document.add_paragraph('')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f'Счет №{dayNumber} от {todayDae}г. на выполнение работ-услуг по договору №{listOfClients.documentNumber} от {listOfClients.documentDate}.')
    run.font.size = Pt(11)
    run.font.bold = True



# <table>--------------------------------------------------------------------------------
    newClArr = []
    newClArr2 = ['', 'Итого:', '', '', '']
    i1 = 0
    for i , l in enumerate(clientArray[1]):
        newClArr.append([])
        newClArr[i].append(str(i+1))
        newClArr[i].append(l[0])
        newClArr[i].append('шт')
        newClArr[i].append(l[1])
        newClArr[i].append(l[2])
        newClArr[i].append(str(int(l[2]) * int(l[1])))
        i1 += int(l[2]) * int(l[1]);
    newClArr2.append(str(i1))
   
    col_names = [
        ['№', 'Наименование', 'Ед.изм.', 'Кол-во', 'Цена, руб.', 'Сумма, руб.'],
        *newClArr,
        newClArr2
    ]

    # tb1 = doc1.add_table(rows=1, cols=len(col_names), style='Colorful Shading Accent 1')
    tb1 = document.add_table(rows=0, cols=0, style = 'Table Grid')
    tb1.add_column(Mm(12.0)) 
    tb1.add_column(Mm(52.0)) 
    tb1.add_column(Mm(17.0)) 
    tb1.add_column(Mm(18.0)) 
    tb1.add_column(Mm(25.0)) 
    tb1.add_column(Mm(29.0)) 

    for d in col_names:
        row = tb1.add_row()   # Add Row Object
        row.height = Mm(8.0)  # Specify row height as 8mm
        
        for i, cell in enumerate(row.cells): # Obtaining Cell object
            cell.text = d[i]                  # Set value to Cell object
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # Set placement position in the cell
# </table>--------------------------------------------------------------------------------


    p = document.add_paragraph(f"Сумма прописью: {num2words(i1, lang='ru')} руб., 00коп. Без НДС.")
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1.5
    p = document.add_paragraph('')



    p = document.add_paragraph()
    run = p.add_run('Руководитель предприятия:')
    run = p.add_run('                                                                               .')
    run.font.underline = True

    p = document.add_paragraph()
    run = p.add_run('Бухгалтер:')
    run = p.add_run('                                                                               .')
    run.font.underline = True
    # run.font.size = Pt(8)



    # document.add_page_break() # Add a page break
    # document.save('demo.docx')
    # https://stackoverflow.com/questions/46011280/how-to-generate-a-docx-in-python-and-save-it-in-memory
    file_stream = io.BytesIO()
    # Save the .docx to the buffer
    document.save(file_stream)
    # Reset the buffer's file-pointer to the beginning of the file
    file_stream.seek(0)

    return file_stream